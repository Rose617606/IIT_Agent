"""批量转换 data/*.docx → data/*.md（YAML frontmatter + Markdown 正文）。

使用方式：
    1. 确保 .env 中已配置 DEEPSEEK_API_KEY
    2. python src/convert_docs.py

流程：
    docx 提取文字（保留标题层级+表格）→ 清洗 → DeepSeek 批量提取元数据 → 生成 .md
"""

import json
import logging
import os
import re
from pathlib import Path

from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn
from openai import OpenAI

load_dotenv()

_logger = logging.getLogger("convert_docs")

_client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url=os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com"),
)
_MODEL = os.getenv("LLM_MODEL", "deepseek-chat")

BATCH_SIZE = 20
DATA_DIR = Path("data")


# ── Step 1: 提取 docx（保留标题 + 表格 + 清洗）────────────

def _detect_heading_level(style_name: str, para) -> int:
    """从段落样式+格式+内容模式判断标题层级（1~3），非标题返回 0。

    三层判断（命中任一层即返回）：
    1. 样式名匹配（Heading 1 / 标题 1 等）
    2. 加粗 + 短段落 → H2
    3. 内容模式匹配（"第X条"→H2，"一、"→H3）
    """
    name = style_name.lower() if style_name else ""

    # 层1：样式名
    for pat, group in [
        (r"heading\s*(\d)", 1), (r"head\s*(\d)", 1),
        (r"标题\s*(\d)", 1), (r"^(\d)\s*级", 1), (r"^title$", 1),
    ]:
        m = re.search(pat, name)
        if m:
            return min(int(m.group(group)), 3)
    if name.strip().isdigit():
        level = int(name.strip())
        if 1 <= level <= 3:
            return level

    # 层2：加粗 + 短段落（政府文件标题不用样式，纯靠加粗）
    is_bold = any(r.bold for r in para.runs if r.bold)
    text = para.text.strip()
    if is_bold and len(text) < 80:
        # 首段加粗 → H1，其余 → H2
        return 1  # 简化处理，加粗短文本即为标题

    # 层3：内容模式（法律/公文特有格式）
    if re.match(r"^第[一二三四五六七八九十百千\d]+[章节条]", text) and len(text) < 100:
        return 2  # "第一条" "第二章" 等
    if re.match(r"^[一二三四五六七八九十]、", text) and len(text) < 100:
        return 3  # "一、..." "二、..."
    if re.match(r"^（[一二三四五六七八九十]）", text) and len(text) < 100:
        return 3  # "（一）..." "（二）..."

    return 0


def _table_to_markdown(table) -> str:
    """将 python-docx Table 转为 Markdown 表格。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        if any(cells):  # 跳过全空行
            rows.append(cells)
    if not rows:
        return ""
    if len(rows) == 1:
        return " | ".join(rows[0])

    lines = []
    # 表头
    lines.append("| " + " | ".join(rows[0]) + " |")
    # 分隔线
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # 数据行
    for row in rows[1:]:
        # 补齐列数（处理合并单元格导致的列数不一致）
        padded = row + [""] * (len(rows[0]) - len(row))
        lines.append("| " + " | ".join(padded[:len(rows[0])]) + " |")
    return "\n".join(lines)


def _clean_text(text: str) -> str:
    """基础清洗：空白符归一、去连续空行、去常见噪声。"""
    # 规范化空白
    text = re.sub(r"[ \t]+", " ", text)           # 多空格 → 单空格
    text = re.sub(r"\n{3,}", "\n\n", text)         # 连续空行 → 最多 2 个
    text = re.sub(r" {2,}\n", "\n", text)          # 行尾多余空格
    text = re.sub(r"\n {2,}", "\n", text)          # 行首多余空格

    # 去纯符号行（页码、分隔线等噪声）
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        # 跳过纯数字行（页码）、纯符号行（装饰线）
        if re.match(r"^[\d\s\-_=*#~•●○·]+$", stripped) and len(stripped) < 8:
            continue
        # 跳过常见页眉页脚
        if re.match(r"^(第\s*\d+\s*页|页码|\d+\s*/\s*\d+)$", stripped):
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def extract_text(docx_path: Path) -> str:
    """从 .docx 提取结构化 Markdown：保留标题层级 + 表格 + 清洗。

    按文档内元素顺序遍历（段落与表格交替），输出格式：
    - 标题段落 → # / ## / ###
    - 普通段落 → 正文
    - 表格 → Markdown 表格
    """
    doc = Document(str(docx_path))
    body_elements: list[str] = []
    p_idx, t_idx = 0, 0

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if p_idx >= len(doc.paragraphs):
                p_idx += 1
                continue
            para = doc.paragraphs[p_idx]
            p_idx += 1
            text = para.text.strip()
            if not text:
                continue
            level = _detect_heading_level(para.style.name if para.style else "", para)
            if level:
                body_elements.append("#" * level + " " + text)
            else:
                body_elements.append(text)

        elif tag == "tbl":
            if t_idx >= len(doc.tables):
                t_idx += 1
                continue
            table = doc.tables[t_idx]
            t_idx += 1
            md_table = _table_to_markdown(table)
            if md_table:
                body_elements.append(md_table)

    raw = "\n\n".join(body_elements)
    return _clean_text(raw)


# ── Step 2: LLM 批量提取元数据 ───────────────────────────

_META_PROMPT = """你是一个中国税法文档处理助手。下面是一批法规文件的「文件名」和「正文开头」，请为每个文件提取元数据，返回 JSON 数组。

严格规则（违反将被拒绝）：
- title: 用文件名，不需要修改
- source: 必须是标准文号（如 "国发〔2018〕41号"、"财税〔2023〕1号"、"国家税务总局公告2020年第1号"、"[20XX]XX号" 等）。没有文号的文件（如问答、解读、案例、提示、通知等非正式文件），source 必须填空字符串 ""，绝对不要填发文机关名、文件名的一部分
- effective_date: 严格按照正文中 "自XXXX年XX月XX日起施行/执行" 提取，格式 YYYY-MM-DD。找不到就填 ""
- status: "active"（现行有效）、"expired"（明确废止/已过有效期）。问答/解读/案例类通常填 "active"

返回格式（严格 JSON 数组，不要 markdown 代码块，不要多余文字）：
[{"filename":"...","title":"...","source":"...","effective_date":"...","status":"..."}]
"""


def extract_metadata_batch(items: list[dict]) -> list[dict]:
    """调用 DeepSeek 批量提取元数据。

    Args:
        items: [{"filename": str, "preview": str}, ...]

    Returns:
        [{"filename": str, "title": str, "source": str,
          "effective_date": str, "status": str}, ...]
    """
    user_content = ""
    for i, item in enumerate(items):
        user_content += f"--- 文件 {i+1} ---\n"
        user_content += f"文件名: {item['filename']}\n"
        user_content += f"正文开头: {item['preview'][:800]}\n\n"

    _logger.info("调用 DeepSeek 提取 %d 个文件的元数据...", len(items))
    resp = _client.chat.completions.create(
        model=_MODEL,
        messages=[
            {"role": "system", "content": _META_PROMPT},
            {"role": "user", "content": user_content},
        ],
        temperature=0.1,
    )
    raw = resp.choices[0].message.content.strip()
    # 去掉可能的 markdown 代码块标记
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        if raw.endswith("```"):
            raw = raw[:-3]
    return json.loads(raw)


# ── Step 3: 生成 .md 文件 ────────────────────────────────

def generate_md(meta: dict, full_text: str, output_path: Path):
    """生成带 YAML frontmatter 的 Markdown 文件。"""
    source = meta.get('source', '') or ''
    effective_date = meta.get('effective_date', '') or '1900-01-01'  # 兜底：未知日期
    yaml_block = f"""---
title: {meta['title']}
source: {source}
effective_date: {effective_date}
status: {meta.get('status', 'active')}
---

"""
    content = yaml_block + full_text
    output_path.write_text(content, encoding="utf-8")


# ── 主入口 ───────────────────────────────────────────────

def convert_all():
    """主流程：扫描 data/*.docx → 提取文字 → LLM 提元数据 → 生成 .md。"""
    docx_files = sorted(DATA_DIR.glob("*.docx"))
    if not docx_files:
        _logger.warning("data/ 目录下没有 .docx 文件")
        return

    _logger.info("共找到 %d 个 .docx 文件", len(docx_files))

    # Step 1: 提取所有文字
    texts: dict[str, str] = {}
    for f in docx_files:
        try:
            texts[f.stem] = extract_text(f)
        except Exception as e:
            _logger.warning("提取失败 %s: %s", f.name, e)

    _logger.info("成功提取 %d 个文件的文字", len(texts))

    # Step 2: 分批 LLM 提取元数据
    items = [
        {"filename": f.stem, "preview": texts[f.stem]}
        for f in docx_files
        if f.stem in texts
    ]

    all_meta: list[dict] = []
    for i in range(0, len(items), BATCH_SIZE):
        batch = items[i : i + BATCH_SIZE]
        try:
            metas = extract_metadata_batch(batch)
            all_meta.extend(metas)
            _logger.info("批次 %d/%d 完成", i // BATCH_SIZE + 1,
                          (len(items) + BATCH_SIZE - 1) // BATCH_SIZE)
        except Exception as e:
            _logger.error("批次 %d 失败: %s", i // BATCH_SIZE + 1, e)
            # 失败的文件用文件名兜底
            for item in batch:
                all_meta.append({
                    "filename": item["filename"],
                    "title": item["filename"],
                    "source": "",
                    "effective_date": "",
                    "status": "active",
                })

    # Step 3: 生成 .md
    meta_by_name = {m["filename"]: m for m in all_meta}
    generated = 0
    for f in docx_files:
        if f.stem not in texts:
            continue
        meta = meta_by_name.get(f.stem, {
            "title": f.stem, "source": "", "effective_date": "", "status": "active"
        })
        md_path = DATA_DIR / f"{f.stem}.md"
        generate_md(meta, texts[f.stem], md_path)
        generated += 1

    _logger.info("完成！共生成 %d 个 .md 文件", generated)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    convert_all()
